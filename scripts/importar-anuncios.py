from __future__ import annotations

import argparse
import csv
import json
import re
import shutil
import sys
import urllib.parse
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.oxml.ns import qn

from anuncios_core import (
    canonical_link,
    check_link,
    classify_product,
    clean_url,
    extract_first_url,
    image_is_placeholder,
    load_json,
    link_is_direct_product,
    normalize_text,
    optimize_image_from_url,
    parse_brazilian_price,
    parse_rating,
    slugify,
    stable_product_id,
    title_similarity,
    today_iso,
    write_json,
)


W_HYPERLINK = qn("w:hyperlink")
R_ID = qn("r:id")
R_EMBED = qn("r:embed")


@dataclass
class ImportedItem:
    title: str
    link: str
    source_file: str
    description: str = ""
    hinted_category: str = ""
    price: str = "Preço no anúncio"
    rating: float | None = None
    embedded_image: bytes | None = None
    embedded_extension: str = ".png"
    embedded_image_safe: bool = False
    source_notes: list[str] = field(default_factory=list)
    validation_link: str = ""
    source_product_link: str = ""
    source: str = "Mercado Livre"
    button_label: str = "Comprar no Mercado Livre"
    product_id_prefix: str = "importado-ml"
    allow_validated_title: bool = False

    def key(self) -> str:
        return canonical_link(self.link) or normalize_text(self.title)


def paragraph_hyperlinks(document: Document, paragraph: Any) -> list[str]:
    result: list[str] = []
    for node in paragraph._p.iter(W_HYPERLINK):
        relationship_id = node.get(R_ID)
        if relationship_id and relationship_id in document.part.rels:
            target = clean_url(document.part.rels[relationship_id].target_ref)
            if target:
                result.append(target)
    return result


def cell_hyperlinks(document: Document, cell: Any) -> list[str]:
    links: list[str] = []
    for paragraph in cell.paragraphs:
        links.extend(paragraph_hyperlinks(document, paragraph))
        links.extend(re.findall(r"https?://[^\s<>]+", paragraph.text))
    return list(dict.fromkeys(clean_url(link) for link in links if clean_url(link)))


def preferred_link(links: Iterable[str], fallback_text: str = "") -> str:
    candidates = list(dict.fromkeys(clean_url(link) for link in links if clean_url(link)))
    candidates.extend(
        link
        for link in re.findall(r"https?://[^\s<>]+", fallback_text or "")
        if clean_url(link) not in candidates
    )
    for candidate in candidates:
        if "meli.la/" in candidate.lower():
            return clean_url(candidate)
    return clean_url(candidates[0]) if candidates else extract_first_url(fallback_text)


def preferred_affiliate_link(links: Iterable[str], fallback_text: str = "") -> str:
    candidates = list(dict.fromkeys(clean_url(link) for link in links if clean_url(link)))
    candidates.extend(
        clean_url(link)
        for link in re.findall(r"https?://[^\s<>]+", fallback_text or "")
        if clean_url(link) not in candidates
    )
    for marker in ("meli.la/", "link.amazon/"):
        for candidate in candidates:
            if marker in candidate.lower():
                return clean_url(candidate)
    for candidate in candidates:
        lowered = candidate.lower()
        if "mercadolivre.com" in lowered or "amazon." in lowered:
            return clean_url(candidate)
    return clean_url(candidates[0]) if candidates else extract_first_url(fallback_text)


def preferred_product_page_link(links: Iterable[str], fallback_text: str = "") -> str:
    candidates = list(dict.fromkeys(clean_url(link) for link in links if clean_url(link)))
    candidates.extend(
        clean_url(link)
        for link in re.findall(r"https?://[^\s<>]+", fallback_text or "")
        if clean_url(link) not in candidates
    )
    direct_markers = (
        "produto.mercadolivre.com",
        "mercadolivre.com.br/",
        "amazon.com.br/dp/",
        "amazon.com.br/gp/product/",
    )
    for candidate in candidates:
        lowered = candidate.lower()
        if any(marker in lowered for marker in direct_markers) and "lista.mercadolivre" not in lowered:
            return clean_url(candidate)
    return preferred_affiliate_link(candidates, fallback_text)


def marketplace_from_link(link: str) -> str:
    lowered = clean_url(link).lower()
    if "amazon." in lowered or "link.amazon" in lowered or "amzn.to" in lowered:
        return "Amazon"
    if "mercadolivre" in lowered or "meli.la" in lowered:
        return "Mercado Livre"
    return "Loja parceira"


def purchase_button_label(link: str) -> str:
    source = marketplace_from_link(link)
    if source == "Amazon":
        return "Comprar na Amazon"
    if source == "Mercado Livre":
        return "Comprar no Mercado Livre"
    return "Comprar na loja parceira"


def cell_image(document: Document, cell: Any) -> tuple[bytes | None, str]:
    for blip in cell._tc.iter(qn("a:blip")):
        relationship_id = blip.get(R_EMBED)
        if not relationship_id or relationship_id not in document.part.rels:
            continue
        part = document.part.rels[relationship_id].target_part
        extension = Path(str(getattr(part, "partname", "image.png"))).suffix or ".png"
        return part.blob, extension
    return None, ".png"


def title_without_number(value: str) -> str:
    return re.sub(r"^\s*\d+\s*[.)-]\s*", "", value or "").strip()


def metadata_value(text: str, label: str) -> str:
    match = re.search(rf"{re.escape(label)}\s*:\s*([^\n]+)", text, re.IGNORECASE)
    return match.group(1).strip() if match else ""


def extract_four_column_catalog(path: Path) -> list[ImportedItem]:
    document = Document(path)
    if not document.tables:
        return []
    table = max(document.tables, key=lambda candidate: len(candidate.rows))
    items: list[ImportedItem] = []
    for row in table.rows[1:]:
        if len(row.cells) < 4:
            continue
        details = row.cells[2].text.strip()
        if not details:
            continue
        title = details.splitlines()[0].strip()
        links = cell_hyperlinks(document, row.cells[3])
        link = preferred_link(links, row.cells[3].text)
        if not title or not link:
            continue
        image, extension = cell_image(document, row.cells[1])
        category = metadata_value(details, "Categoria")
        items.append(
            ImportedItem(
                title=title,
                link=link,
                source_file=path.name,
                description=details,
                hinted_category=category,
                price=parse_brazilian_price(details),
                rating=parse_rating(details),
                embedded_image=image,
                embedded_extension=extension,
                embedded_image_safe=bool(image),
                source_notes=["Foto incorporada ao lado do produto no catálogo."],
            )
        )
    return items


def extract_lanterns(path: Path) -> list[ImportedItem]:
    document = Document(path)
    items: list[ImportedItem] = []
    for table in document.tables:
        if len(table.rows) != 1 or len(table.columns) != 2:
            continue
        image_cell, detail_cell = table.rows[0].cells
        text = detail_cell.text.strip()
        if not re.match(r"^\d+\.", text):
            continue
        title = title_without_number(text.splitlines()[0])
        links = cell_hyperlinks(document, detail_cell)
        link = preferred_link(links, text)
        if not link:
            continue
        image, extension = cell_image(document, image_cell)
        items.append(
            ImportedItem(
                title=title,
                link=link,
                source_file=path.name,
                description=text,
                hinted_category="Lanternas UV",
                price=parse_brazilian_price(text),
                rating=parse_rating(text),
                embedded_image=image,
                embedded_extension=extension,
                embedded_image_safe=bool(image),
                source_notes=["Foto incorporada no bloco individual da lanterna."],
            )
        )
    return items[:10]


def extract_tvs(path: Path) -> list[ImportedItem]:
    document = Document(path)
    numbered_titles = [
        title_without_number(paragraph.text)
        for paragraph in document.paragraphs
        if re.match(r"^\s*\d+\.", paragraph.text)
    ]
    items: list[ImportedItem] = []
    for index, table in enumerate(document.tables):
        if index >= len(numbered_titles) or not table.rows or len(table.rows[0].cells) < 2:
            continue
        image_cell, detail_cell = table.rows[0].cells
        detail = detail_cell.text.strip()
        links = cell_hyperlinks(document, detail_cell)
        link = preferred_link(links, detail)
        if not link:
            continue
        image, extension = cell_image(document, image_cell)
        items.append(
            ImportedItem(
                title=numbered_titles[index],
                link=link,
                source_file=path.name,
                description=detail,
                hinted_category="Smart TVs",
                price="Preço no anúncio",
                embedded_image=image,
                embedded_extension=extension,
                embedded_image_safe=False,
                source_notes=[
                    "O próprio arquivo classifica as fotos como ilustrativas/representativas; não publicar sem validação online."
                ],
            )
        )
    return items


def row_hyperlinks(document: Document, row: Any) -> list[str]:
    links: list[str] = []
    for cell in row.cells:
        links.extend(cell_hyperlinks(document, cell))
    return list(dict.fromkeys(clean_url(link) for link in links if clean_url(link)))


def choose_direct_product_link(links: Iterable[str]) -> str:
    candidates = list(dict.fromkeys(clean_url(link) for link in links if clean_url(link)))
    for candidate in candidates:
        if link_is_direct_product(candidate) and "meli.la/" not in candidate.lower():
            return candidate
    for candidate in candidates:
        if "meli.la/" in candidate.lower():
            return candidate
    return candidates[0] if candidates else ""


def clean_catalog_text(value: str) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    return (
        text.replace("Cristă", "Cristã")
        .replace("cristă", "cristã")
        .replace("Biblia", "Bíblia")
    )


def clean_amazon_title(value: str) -> str:
    title = re.sub(r"\s+", " ", str(value or "")).strip()
    title = re.sub(r"\s*\|\s*Amazon(?:\.com\.br)?\s*$", "", title, flags=re.IGNORECASE).strip()
    return title or str(value or "").strip()


def amazon_link_kind(value: str) -> str:
    value = clean_url(value)
    if not value:
        return ""
    try:
        parsed = urllib.parse.urlsplit(value)
    except ValueError:
        return ""
    host = parsed.netloc.lower().removeprefix("www.")
    if host in {"link.amazon", "amzn.to"}:
        return "affiliate"
    if "amazon." in host:
        return "product" if link_is_direct_product(value) else "search"
    return ""


def split_amazon_links(links: Iterable[str]) -> tuple[str, str, str]:
    clean_links = list(dict.fromkeys(clean_url(link) for link in links if clean_url(link)))
    affiliate = next((link for link in clean_links if amazon_link_kind(link) == "affiliate"), "")
    product = next((link for link in clean_links if amazon_link_kind(link) == "product"), "")
    search = next((link for link in clean_links if amazon_link_kind(link) == "search"), "")
    return affiliate, product, search


def amazon_item_from_parts(
    *,
    title: str,
    link: str,
    validation_link: str,
    source_product_link: str,
    source_file: str,
    description: str,
    hinted_category: str,
    price: str = "Preço no anúncio",
    rating: float | None = None,
    embedded_image: bytes | None = None,
    embedded_extension: str = ".png",
    embedded_image_safe: bool = False,
    source_notes: list[str] | None = None,
    allow_validated_title: bool = False,
) -> ImportedItem:
    return ImportedItem(
        title=title,
        link=link,
        source_file=source_file,
        description=description,
        hinted_category=hinted_category,
        price=price,
        rating=rating,
        embedded_image=embedded_image,
        embedded_extension=embedded_extension,
        embedded_image_safe=embedded_image_safe,
        source_notes=source_notes or [],
        validation_link=validation_link or link,
        source_product_link=source_product_link,
        source="Amazon",
        button_label="Comprar na Amazon",
        product_id_prefix="importado-amazon",
        allow_validated_title=allow_validated_title,
    )


def extract_amazon_card_catalog(path: Path) -> list[ImportedItem]:
    document = Document(path)
    items: list[ImportedItem] = []
    for table in document.tables:
        if not table.rows or len(table.rows[0].cells) < 2:
            continue
        image_cell, detail_cell = table.rows[0].cells[0], table.rows[0].cells[1]
        raw_text = detail_cell.text.strip()
        if not re.match(r"^\s*\d{1,3}\.", raw_text):
            continue
        lines = [clean_catalog_text(line) for line in raw_text.splitlines() if clean_catalog_text(line)]
        if not lines:
            continue
        title = title_without_number(lines[0])
        if not title:
            continue
        links = cell_hyperlinks(document, detail_cell)
        affiliate, product_link, search_link = split_amazon_links(links)
        link = affiliate
        validation_link = product_link or affiliate
        if not link or not validation_link:
            continue
        image, extension = cell_image(document, image_cell)
        category = metadata_value(raw_text, "Categoria") or "Amazon"
        description = " ".join(lines)
        notes = [
            "Fonte: catálogo Amazon enviado pelo usuário.",
            "Foto e preço devem ser validados pelo link direto da Amazon; botão da loja usa somente o link de afiliado.",
        ]
        if search_link and not product_link:
            notes.append("O arquivo também traz link de busca/conferência; item depende do link de afiliado para validar produto exato.")
        items.append(
            amazon_item_from_parts(
                title=title,
                link=link,
                validation_link=validation_link,
                source_product_link=product_link or search_link,
                source_file=path.name,
                description=description,
                hinted_category=category,
                price=parse_brazilian_price(raw_text),
                rating=parse_rating(raw_text),
                embedded_image=image,
                embedded_extension=extension,
                embedded_image_safe=False,
                source_notes=notes,
            )
        )
    return items


def extract_amazon_trend_tables(path: Path) -> list[ImportedItem]:
    document = Document(path)
    items: list[ImportedItem] = []
    for table in document.tables:
        if len(table.rows) < 2:
            continue
        headers = [normalize_text(cell.text) for cell in table.rows[0].cells]
        if "produto termo" not in " ".join(headers) or "link de afiliado" not in " ".join(headers):
            continue
        for row in table.rows[1:]:
            cells = row.cells
            if len(cells) < 10:
                continue
            title = clean_catalog_text(cells[1].text)
            if not title:
                continue
            category = clean_catalog_text(cells[2].text) or "Amazon"
            direct_links = cell_hyperlinks(document, cells[8])
            affiliate_links = cell_hyperlinks(document, cells[9])
            affiliate, product_link, search_link = split_amazon_links(direct_links + affiliate_links)
            link = affiliate
            validation_link = product_link or affiliate
            if not link or not validation_link:
                continue
            rank = clean_catalog_text(cells[0].text)
            commission = clean_catalog_text(cells[3].text)
            score = clean_catalog_text(cells[4].text)
            trend = clean_catalog_text(cells[5].text)
            opportunity = clean_catalog_text(cells[6].text)
            dates = clean_catalog_text(cells[7].text)
            description = (
                f"Produto/termo {rank} do catálogo de tendências Amazon. "
                f"Categoria informada: {category}. Comissão estimada: {commission}. "
                f"Score: {score}. Tendência: {trend}. Oportunidade: {opportunity}. "
                f"Datas fortes: {dates}. Confira preço, variação, vendedor, frete e disponibilidade diretamente na Amazon."
            )
            notes = [
                "Fonte: catálogo de produtos buscados no Brasil e tendências para afiliado.",
                "O link de busca foi usado apenas como referência; o botão da loja preserva o link de afiliado.",
            ]
            if search_link:
                notes.append("Produto veio de termo de busca; quando o link afiliado resolve para item exato, o título validado da Amazon é priorizado.")
            items.append(
                amazon_item_from_parts(
                    title=title,
                    link=link,
                    validation_link=validation_link,
                    source_product_link=product_link or search_link,
                    source_file=path.name,
                    description=description,
                    hinted_category=category,
                    source_notes=notes,
                    allow_validated_title=True,
                )
            )
    return items


def extract_amazon_ovens_microwaves(path: Path) -> list[ImportedItem]:
    document = Document(path)
    items: list[ImportedItem] = []
    for table_index, table in enumerate(document.tables):
        if len(table.rows) < 2 or len(table.rows[0].cells) < 9:
            continue
        headers = [normalize_text(cell.text) for cell in table.rows[0].cells]
        if "link direto amazon" not in " ".join(headers) or "link de afiliado" not in " ".join(headers):
            continue
        group = "Fornos de mesa" if table_index == 0 else "Micro-ondas"
        for row in table.rows[1:]:
            cells = row.cells
            title = clean_catalog_text(cells[1].text)
            if not title:
                continue
            voltage = clean_catalog_text(cells[3].text)
            if voltage and voltage not in title:
                title = f"{title} - {voltage}"
            links = []
            links.extend(cell_hyperlinks(document, cells[7]))
            links.extend(cell_hyperlinks(document, cells[8]))
            affiliate, product_link, search_link = split_amazon_links(links)
            link = affiliate
            validation_link = product_link or affiliate
            if not link or not validation_link:
                continue
            description = (
                f"Item {clean_catalog_text(cells[0].text)} do catálogo Amazon de {group}. "
                f"Capacidade: {clean_catalog_text(cells[2].text)}. "
                f"Voltagem: {clean_catalog_text(cells[3].text)}. "
                f"Avaliação informada: {clean_catalog_text(cells[4].text)}. "
                f"Envio/entrega: {clean_catalog_text(cells[5].text)}. "
                f"Comissão estimada: {clean_catalog_text(cells[6].text)}. "
                "Confira preço, voltagem, frete e disponibilidade diretamente na Amazon."
            )
            items.append(
                amazon_item_from_parts(
                    title=title,
                    link=link,
                    validation_link=validation_link,
                    source_product_link=product_link or search_link,
                    source_file=path.name,
                    description=description,
                    hinted_category=group,
                    rating=parse_rating(cells[4].text),
                    source_notes=[
                        "Fonte: lista corrigida de fornos de mesa e micro-ondas Amazon Brasil.",
                        "Foto e preço importados a partir da página original; botão usa link de afiliado.",
                    ],
                )
            )
    return items


def extract_amazon_catalog(path: Path) -> list[ImportedItem]:
    name = normalize_text(path.name)
    if "buscados brasil tendencias" in name:
        return extract_amazon_trend_tables(path)
    if "fornos de mesa microondas" in name:
        return extract_amazon_ovens_microwaves(path)
    return extract_amazon_card_catalog(path)


def extract_ovens(path: Path) -> list[ImportedItem]:
    document = Document(path)
    data_tables = [
        table
        for table in document.tables
        if len(table.rows) > 1 and len(table.columns) >= 7
    ]
    if not data_tables:
        return []
    table = max(data_tables, key=lambda candidate: len(candidate.rows))
    items: list[ImportedItem] = []
    for row in table.rows[1:]:
        cells = row.cells
        number = cells[0].text.strip()
        title = clean_catalog_text(cells[1].text)
        if not title:
            continue
        links = row_hyperlinks(document, row)
        link = choose_direct_product_link(links)
        if not link:
            continue
        liters = clean_catalog_text(cells[2].text)
        reason = clean_catalog_text(cells[3].text)
        indicated_for = clean_catalog_text(cells[4].text)
        if "http" in indicated_for.lower():
            indicated_for = ""
        price_text = clean_catalog_text(cells[5].text)
        description_parts = [
            f"Item {number} do catálogo de fornos de bancada.",
            f"Capacidade: {liters}." if liters else "",
            reason,
            f"Indicado para: {indicated_for}." if indicated_for else "",
            price_text,
        ]
        short_links = [candidate for candidate in links if "meli.la/" in candidate.lower()]
        direct_links = [candidate for candidate in links if "mercadolivre.com" in candidate.lower()]
        source_notes = ["Fonte: curadoria de fornos de bancada Mercado Livre."]
        if short_links and direct_links and link not in short_links:
            source_notes.append(
                "Link direto do produto priorizado porque o link curto do arquivo pode apontar para lista ou recomendação."
            )
        items.append(
            ImportedItem(
                title=title,
                link=link,
                source_file=path.name,
                description=" ".join(part for part in description_parts if part),
                hinted_category="Fornos elétricos de bancada",
                price=parse_brazilian_price(price_text),
                rating=parse_rating(reason),
                embedded_image_safe=False,
                source_notes=source_notes,
            )
        )
    return items


def extract_evangelical(path: Path) -> list[ImportedItem]:
    document = Document(path)
    items: list[ImportedItem] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if not re.search(r"\bITEM\s+\d+\b", text, re.IGNORECASE):
                    continue
                blocks = re.split(r"(?=\bITEM\s+\d+\b)", text, flags=re.IGNORECASE)
                cell_links = cell_hyperlinks(document, cell)
                for block in blocks:
                    if not re.search(r"\bITEM\s+\d+\b", block, re.IGNORECASE):
                        continue
                    lines = [
                        clean_catalog_text(line)
                        for line in block.splitlines()
                        if clean_catalog_text(line)
                    ]
                    if len(lines) < 2:
                        continue
                    item_number = lines[0]
                    title = lines[1]
                    category = next(
                        (
                            line
                            for line in lines[2:]
                            if "mercado livre" not in normalize_text(line)
                            and not line.lower().startswith("http")
                            and not line.lower().startswith("link direto")
                            and not line.lower().startswith("busca alternativa")
                        ),
                        "Produtos religiosos",
                    )
                    block_links = re.findall(r"https?://[^\s<>]+", block)
                    links = list(dict.fromkeys(block_links + cell_links))
                    direct_links = [
                        link for link in links if link_is_direct_product(link) and "meli.la/" in link.lower()
                    ]
                    link = direct_links[0] if direct_links else preferred_link(links, block)
                    if not title or not link:
                        continue
                    items.append(
                        ImportedItem(
                            title=title,
                            link=link,
                            source_file=path.name,
                            description=(
                                f"{item_number} do catálogo evangélico. "
                                f"Categoria informada: {category}. "
                                "Confira preço, edição, acabamento, frete e disponibilidade diretamente no anúncio."
                            ),
                            hinted_category=category,
                            price="Preço no anúncio",
                            embedded_image_safe=False,
                            source_notes=[
                                "Fonte: catálogo evangélico com links atualizados do Mercado Livre."
                            ],
                        )
                    )
    return items


def extract_cookware_catalog(path: Path) -> list[ImportedItem]:
    document = Document(path)
    items: list[ImportedItem] = []
    catalog_note = " ".join(
        clean_catalog_text(paragraph.text)
        for paragraph in document.paragraphs
        if clean_catalog_text(paragraph.text)
    )
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if not re.search(r"\bITEM\s+\d+\b", text, re.IGNORECASE):
                    continue
                blocks = re.split(r"(?=\bITEM\s+\d+\b)", text, flags=re.IGNORECASE)
                cell_links = cell_hyperlinks(document, cell)
                for block in blocks:
                    if not re.search(r"\bITEM\s+\d+\b", block, re.IGNORECASE):
                        continue
                    lines = [
                        clean_catalog_text(line)
                        for line in block.splitlines()
                        if clean_catalog_text(line)
                    ]
                    if len(lines) < 2:
                        continue
                    item_number = lines[0]
                    title = lines[1]
                    block_links = re.findall(r"https?://[^\s<>]+", block)
                    links = list(dict.fromkeys(block_links + cell_links))
                    short_links = [
                        candidate for candidate in links if "meli.la/" in candidate.lower()
                    ]
                    link = short_links[0] if short_links else choose_direct_product_link(links)
                    if not title or not link:
                        continue
                    search_only = not link_is_direct_product(link)
                    notes = [
                        "Fonte: catálogo de jogos de panelas Pacific e similares com links Mercado Livre."
                    ]
                    if search_only:
                        notes.append(
                            "Item com link de busca/lista; manter em revisão manual até confirmar o anúncio exato."
                        )
                    description = (
                        f"{item_number} do catálogo de panelas Pacific e similares. "
                        "Categoria informada: Jogos de panelas. "
                        "Confira preço, quantidade de peças, compatibilidade com indução, frete e disponibilidade diretamente no anúncio."
                    )
                    if catalog_note:
                        description = f"{description} Observação do arquivo: {catalog_note}"
                    items.append(
                        ImportedItem(
                            title=title,
                            link=link,
                            source_file=path.name,
                            description=description,
                            hinted_category="Panelas",
                            price="Preço no anúncio",
                            embedded_image_safe=False,
                            source_notes=notes,
                        )
                    )
    return items


def extract_battery_toys(path: Path) -> list[ImportedItem]:
    document = Document(path)
    quick_table = max(document.tables, key=lambda candidate: len(candidate.rows))
    short_links_by_number: dict[str, str] = {}
    for table in document.tables:
        if table._tbl is quick_table._tbl:
            continue
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                number_match = re.match(r"^\s*(\d+)\.", text)
                if not number_match:
                    continue
                link = preferred_link(cell_hyperlinks(document, cell), text)
                if "meli.la/" in link.lower():
                    short_links_by_number[number_match.group(1)] = link
    items: list[ImportedItem] = []
    for row in quick_table.rows[1:]:
        if len(row.cells) < 3:
            continue
        number = row.cells[0].text.strip()
        title = row.cells[1].text.strip()
        links = cell_hyperlinks(document, row.cells[2])
        link = short_links_by_number.get(number) or preferred_link(links, row.cells[2].text)
        if not title or not link:
            continue
        items.append(
            ImportedItem(
                title=title,
                link=link,
                source_file=path.name,
                description="Brinquedo a bateria/recarregável selecionado no catálogo enviado.",
                hinted_category="Brinquedos a bateria",
                price="Preço no anúncio",
                embedded_image_safe=False,
                source_notes=[
                    "O próprio arquivo informa que as imagens são ilustrações por categoria; imagem real será buscada no anúncio."
                ],
            )
        )
    return items


def extract_offer_tables(path: Path) -> list[ImportedItem]:
    document = Document(path)
    items: list[ImportedItem] = []
    normalized_name = normalize_text(path.name)
    default_category = "Materiais escolares" if "materiais" in normalized_name else "Calcados"
    for table_index, table in enumerate(document.tables):
        if not table.rows:
            continue
        headers = [normalize_text(cell.text) for cell in table.rows[0].cells]
        if "produto" not in headers or "categoria" not in headers:
            continue
        product_index = headers.index("produto")
        category_index = headers.index("categoria")
        brand_index = headers.index("marca") if "marca" in headers else None
        for row in table.rows[1:]:
            cells = list(row.cells)
            if product_index >= len(cells):
                continue
            title = title_without_number(cells[product_index].text.strip())
            if not title:
                continue
            row_text = " ".join(cell.text for cell in cells)
            row_links: list[str] = []
            for cell in cells:
                row_links.extend(cell_hyperlinks(document, cell))
            link = preferred_affiliate_link(row_links, row_text)
            if not link:
                continue
            product_page_link = preferred_product_page_link(row_links, row_text)
            category = cells[category_index].text.strip() if category_index < len(cells) else ""
            brand = cells[brand_index].text.strip() if brand_index is not None and brand_index < len(cells) else ""
            if "materiais" in normalized_name:
                hinted_category = f"loja livraria prioridade {default_category}"
            elif "calcados" in normalized_name:
                hinted_category = f"loja calcados prioridade {category or default_category}"
            else:
                hinted_category = category or default_category
            details = [title]
            if category:
                details.append(f"Categoria: {category}")
            if brand:
                details.append(f"Marca: {brand}")
            if "materiais" in normalized_name:
                details.append("Produto selecionado na lista de materiais escolares enviada para a livraria.")
            elif "calcados" in normalized_name:
                details.append("Produto selecionado na lista de calcados enviada para a loja de calcados.")
            else:
                details.append("Produto selecionado no catalogo enviado.")
            details.append("Confira preco, estoque, frete, garantia, tamanho/cor e condicoes diretamente no anuncio.")
            source = marketplace_from_link(link or product_page_link)
            section = "Mercado Livre" if "meli.la/" in link.lower() else source
            items.append(
                ImportedItem(
                    title=title,
                    link=link,
                    source_file=f"{path.name} | Tabela {table_index + 1} {section}",
                    description=". ".join(part for part in details if part),
                    hinted_category=hinted_category,
                    price=parse_brazilian_price(row_text),
                    embedded_image_safe=False,
                    source_notes=[
                        "Foto buscada a partir do link do produto; link de afiliado preservado no cadastro."
                    ],
                    validation_link=product_page_link or link,
                    source_product_link=product_page_link,
                    source=source,
                    button_label=purchase_button_label(link),
                    product_id_prefix="importado-amazon" if source == "Amazon" else "importado-ml",
                    allow_validated_title=False,
                )
            )
    return items


def extract_docx(path: Path) -> list[ImportedItem]:
    name = normalize_text(path.name)
    if "materiais escolares" in name or "calcados" in name:
        return extract_offer_tables(path)
    if "produtos buscados" in name and "tendencias" in name:
        return extract_amazon_trend_tables(path)
    if "amazon" in name:
        return extract_amazon_catalog(path)
    if "panelas" in name or "pacific" in name:
        return extract_cookware_catalog(path)
    if "fornos bancada" in name or "forno" in name:
        return extract_ovens(path)
    if "evangelico" in name or "evangelicos" in name:
        return extract_evangelical(path)
    if "lanternas uv" in name:
        return extract_lanterns(path)
    if "tvs mercado livre" in name:
        return extract_tvs(path)
    if "50 brinquedos a bateria" in name:
        return extract_battery_toys(path)
    return extract_four_column_catalog(path)


def extract_json(path: Path) -> list[ImportedItem]:
    payload = json.loads(path.read_text(encoding="utf-8"))
    records = payload if isinstance(payload, list) else payload.get("products", payload.get("produtos", []))
    result = []
    for record in records:
        title = str(record.get("name") or record.get("nome") or record.get("title") or "").strip()
        link = str(
            record.get("affiliateLink")
            or record.get("linkCompra")
            or record.get("link")
            or record.get("url")
            or ""
        ).strip()
        if title and link:
            result.append(
                ImportedItem(
                    title=title,
                    link=link,
                    source_file=path.name,
                    description=str(record.get("description") or record.get("descricao") or ""),
                    hinted_category=str(record.get("category") or record.get("categoria") or ""),
                    price=str(record.get("price") or record.get("preco") or "Preço no anúncio"),
                )
            )
    return result


def extract_csv_file(path: Path) -> list[ImportedItem]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        records = list(csv.DictReader(handle))
    temporary = path.with_suffix(".json")
    temporary.write_text(json.dumps(records, ensure_ascii=False), encoding="utf-8")
    try:
        return extract_json(temporary)
    finally:
        temporary.unlink(missing_ok=True)


def extract_plain_text(path: Path) -> list[ImportedItem]:
    result = []
    for line in path.read_text(encoding="utf-8", errors="replace").splitlines():
        link = extract_first_url(line)
        title = re.sub(r"https?://[^\s<>]+", "", line).strip(" -|;\t")
        if title and link:
            result.append(ImportedItem(title=title, link=link, source_file=path.name))
    return result


def extract_file(path: Path) -> list[ImportedItem]:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".json":
        return extract_json(path)
    if suffix in {".csv", ".tsv"}:
        return extract_csv_file(path)
    if suffix in {".txt", ".html", ".htm"}:
        return extract_plain_text(path)
    raise ValueError(f"Formato não suportado: {path.name}")


def richness(item: ImportedItem) -> int:
    return (
        len(item.description)
        + (1000 if item.embedded_image_safe and item.embedded_image else 0)
        + (200 if item.rating is not None else 0)
        + (100 if item.price != "Preço no anúncio" else 0)
    )


def merge_items(items: Iterable[ImportedItem]) -> list[ImportedItem]:
    merged: dict[str, ImportedItem] = {}
    for item in items:
        key = item.key()
        if not key:
            continue
        if key not in merged:
            merged[key] = item
            continue
        current = merged[key]
        preferred, other = (item, current) if richness(item) > richness(current) else (current, item)
        preferred.source_notes = list(dict.fromkeys(preferred.source_notes + other.source_notes))
        preferred.source_file = " | ".join(
            dict.fromkeys(part.strip() for part in f"{preferred.source_file}|{other.source_file}".split("|"))
        )
        if not preferred.embedded_image and other.embedded_image:
            preferred.embedded_image = other.embedded_image
            preferred.embedded_extension = other.embedded_extension
            preferred.embedded_image_safe = other.embedded_image_safe
        merged[key] = preferred
    return list(merged.values())


def ensure_toy_store(stores: list[dict[str, Any]]) -> bool:
    if any(store.get("id") == "impacto-brinquedos" for store in stores):
        return False
    stores.append(
        {
            "id": "impacto-brinquedos",
            "slug": "brinquedos",
            "route": "/loja/brinquedos",
            "name": "IMPACTO BRINQUEDOS",
            "commercialName": "Brinquedos",
            "floor": "segundo-andar",
            "category": "Brinquedos",
            "type": "products",
            "theme": "toys",
            "icon": "TOY",
            "color": "#ee7b2d",
            "gradient": "linear-gradient(135deg,#fff5df,#ffffff 46%,#e7f7ff)",
            "description": (
                "Brinquedos selecionados, ofertas infantis, brinquedos educativos, "
                "carrinhos, bonecas, drones, helicópteros e brinquedos a bateria."
            ),
            "subcategories": [
                "Brinquedos educativos",
                "Bonecas",
                "Carrinhos",
                "Drones infantis",
                "Helicópteros",
                "Jogos",
                "Blocos de montar",
                "Brinquedos a bateria",
                "Faz de conta",
                "Robôs e dinossauros",
            ],
            "section": "Brincar, aprender e imaginar",
            "active": True,
        }
    )
    return True


def existing_index(products: list[dict[str, Any]]) -> tuple[dict[str, dict[str, Any]], dict[str, dict[str, Any]]]:
    by_link: dict[str, dict[str, Any]] = {}
    by_title: dict[str, dict[str, Any]] = {}
    for product in products:
        link = canonical_link(
            product.get("affiliateLink")
            or product.get("linkCompra")
            or product.get("linkOriginal")
            or product.get("linkPrincipalFonte")
            or ""
        )
        if link:
            by_link.setdefault(link, product)
        title = normalize_text(product.get("name") or product.get("nome") or "")
        if title:
            by_title.setdefault(title, product)
    return by_link, by_title


def build_product(
    item: ImportedItem,
    check: dict[str, Any],
    root: Path,
    image_relative: str,
    existing: dict[str, Any] | None = None,
) -> dict[str, Any]:
    product = dict(existing or {})
    validation_link = item.validation_link or item.link
    checked_title = clean_amazon_title(check.get("title", "")) if item.source == "Amazon" else check.get("title", "")
    display_title = (
        checked_title
        if item.allow_validated_title and check.get("ok") and checked_title
        else item.title
    )
    classification = classify_product(display_title, item.description, item.hinted_category)
    match = title_similarity(display_title, checked_title or check.get("title", ""))
    validated_link_is_product = (
        link_is_direct_product(validation_link)
        or link_is_direct_product(check.get("finalUrl", ""))
        or link_is_direct_product(check.get("directProductUrl", ""))
    )
    valid_online = (
        validated_link_is_product
        and bool(check.get("ok"))
        and bool(match["safe"])
    )
    status = "ativo" if valid_online and image_relative else "revisao_manual"
    raw_price = check.get("price") or item.price or product.get("price") or "Preço no anúncio"
    price = parse_brazilian_price(raw_price)
    rating = check.get("rating") if check.get("rating") is not None else item.rating
    product_id = product.get("id") or stable_product_id(
        item.product_id_prefix, display_title, item.link
    )
    description = item.description.strip() or (
        f"Produto selecionado no arquivo {item.source_file}. "
        "Confira preço, frete, garantia e condições diretamente no anúncio."
    )
    pendencias: list[str] = []
    if not check.get("ok"):
        pendencias.append("link não confirmado online")
    if not validated_link_is_product:
        pendencias.append("link de produto não confirmado")
    if check.get("title") and not match["safe"]:
        pendencias.append("título do link diverge do produto")
    if not image_relative:
        pendencias.append("imagem real não confirmada")
    if price == "Preço no anúncio":
        pendencias.append("preço dinâmico no anúncio")

    product.update(
        {
            "id": product_id,
            "storeId": classification["storeId"],
            "name": display_title,
            "description": description,
            "descricaoCurta": description,
            "descricaoDetalhada": description,
            "price": price,
            "image": image_relative or "public/placeholder-produto-mercado-livre.svg",
            "imagemPrincipal": image_relative or "public/placeholder-produto-mercado-livre.svg",
            "galeria": [image_relative] if image_relative else [],
            "badge": "Oferta verificada" if status == "ativo" else "Revisar",
            "affiliateLink": item.link,
            "linkCompra": item.link,
            "linkOriginal": item.link,
            "linkResolvidoApenasLeitura": check.get("directProductUrl") or check.get("finalUrl") or item.source_product_link or "",
            "linkProdutoApenasLeitura": item.source_product_link or check.get("finalUrl") or "",
            "category": classification["category"],
            "categoria": classification["category"],
            "subcategoria": classification["subcategoria"],
            "source": item.source,
            "origem": item.source,
            "status": status,
            "aprovadoParaPublicacao": status == "ativo",
            "destaqueHome": True,
            "homeRotation": True,
            "rotacaoTelaInicial": True,
            "editable": True,
            "actionType": "buy",
            "buttonLabel": item.button_label,
            "rating": rating,
            "nota": rating,
            "origemImportacao": item.source_file,
            "arquivosOrigem": [part.strip() for part in item.source_file.split("|")],
            "statusImagem": "imagem real confirmada" if image_relative else "imagem manual necessária",
            "linkStatus": "link confirmado" if valid_online else "revisão manual",
            "pendencias": pendencias,
            "observacaoImportacao": " ".join(item.source_notes),
            "validacaoOnline": {
                **check,
                "titleMatch": match,
                "affiliateLinkCadastrado": item.link,
                "validationLinkUsadoParaFoto": validation_link,
                "sourceProductLink": item.source_product_link,
                "checkedAt": today_iso(),
            },
            "ultimaRevisao": today_iso(),
        }
    )
    return product


def product_link(product: dict[str, Any]) -> str:
    return str(
        product.get("affiliateLink")
        or product.get("linkCompra")
        or product.get("linkOriginal")
        or ""
    ).strip()


def product_image(product: dict[str, Any]) -> str:
    return str(product.get("image") or product.get("imagemPrincipal") or "").strip()


def product_summary(product: dict[str, Any]) -> str:
    description = str(
        product.get("descricaoCurta") or product.get("description") or product.get("name") or ""
    ).strip()
    description = re.sub(r"\s+", " ", description)
    if len(description) > 170:
        description = description[:167].rstrip() + "..."
    return description


def update_home_rotation(root: Path, imported_products: list[dict[str, Any]]) -> dict[str, int]:
    banners_path = root / "dados" / "banners-anuncios.json"
    data = load_json(
        banners_path,
        {"settings": {"bannerRotationMs": 6500, "adRotationMs": 5200}, "banners": [], "ads": []},
    )
    data["settings"] = data.get("settings") or {"bannerRotationMs": 6500, "adRotationMs": 5200}
    data["banners"] = data.get("banners") if isinstance(data.get("banners"), list) else []
    data["ads"] = data.get("ads") if isinstance(data.get("ads"), list) else []
    ads: list[dict[str, Any]] = data["ads"]
    by_id = {str(ad.get("id")): ad for ad in ads}
    priority = max((int(ad.get("priority") or 0) for ad in ads), default=0)
    prepared = 0
    active = 0
    for product in imported_products:
        ad_id = f"ad-produto-{product.get('id')}"
        link = product_link(product)
        image = product_image(product)
        ready = (
            str(product.get("status", "")).lower() == "ativo"
            and bool(link)
            and bool(image)
            and not image_is_placeholder(image)
        )
        prepared += 1
        if ready:
            active += 1
        product["homeRotation"] = True
        product["rotacaoTelaInicial"] = True
        product["destaqueHome"] = True
        product["homeRotationAdId"] = ad_id
        current = by_id.get(ad_id)
        if current is None:
            priority += 1
            current = {"id": ad_id, "priority": priority}
            ads.append(current)
        current.update(
            {
                "productId": product.get("id"),
                "storeId": product.get("storeId"),
                "image": image,
                "title": product.get("name") or product.get("nome") or "",
                "description": product_summary(product),
                "buttonLabel": product.get("buttonLabel") or "Ver oferta",
                "link": link,
                "startDate": today_iso(),
                "endDate": "",
                "active": ready,
                "source": "produto-importado-docx",
                "rotationGroup": "materiais-calcados-2026-07-09",
            }
        )
    write_json(banners_path, data)
    return {"prepared": prepared, "active": active, "inactive": prepared - active}


def check_item_link(item: ImportedItem) -> dict[str, Any]:
    candidates = [item.link]
    for candidate in (item.validation_link, item.source_product_link):
        if candidate and canonical_link(candidate) != canonical_link(item.link):
            candidates.append(candidate)
    last_result: dict[str, Any] | None = None
    for candidate in list(dict.fromkeys(candidates)):
        result = check_link(candidate, item.title).as_dict()
        if item.source_product_link:
            result["sourceProductLink"] = item.source_product_link
        if candidate != item.link:
            result["validationLinkFallback"] = candidate
        if result.get("ok"):
            return result
        last_result = result
    return last_result or check_link(item.link, item.title).as_dict()


def main() -> int:
    parser = argparse.ArgumentParser(description="Importa anúncios e aplica filtros de qualidade.")
    parser.add_argument("inputs", nargs="+", type=Path)
    parser.add_argument("--root", type=Path, default=Path.cwd())
    parser.add_argument("--apply", action="store_true")
    parser.add_argument("--offline", action="store_true")
    parser.add_argument("--workers", type=int, default=6)
    args = parser.parse_args()

    root = args.root.resolve()
    products_path = root / "dados" / "products.json"
    stores_path = root / "dados" / "stores.json"
    products: list[dict[str, Any]] = load_json(products_path, [])
    stores: list[dict[str, Any]] = load_json(stores_path, [])

    extracted: list[ImportedItem] = []
    extraction_counts: dict[str, int] = {}
    for input_path in args.inputs:
        resolved = input_path.resolve()
        items = extract_file(resolved)
        extracted.extend(items)
        extraction_counts[resolved.name] = len(items)
    merged = merge_items(extracted)

    checks: dict[str, dict[str, Any]] = {}
    if args.offline:
        for item in merged:
            checks[item.key()] = {
                "requestedUrl": item.link,
                "ok": True,
                "title": item.title,
                "image": "",
                "error": "validação online ignorada",
            }
    else:
        with ThreadPoolExecutor(max_workers=max(1, args.workers)) as executor:
            pending = {
                executor.submit(check_item_link, item): item
                for item in merged
            }
            for future in as_completed(pending):
                item = pending[future]
                checks[item.key()] = future.result()
                status = "OK" if checks[item.key()].get("ok") else "REVISAR"
                print(f"[{status}] {item.title[:72]}")

    selected_by_title: dict[str, ImportedItem] = {}
    for item in merged:
        title_key = normalize_text(item.title)
        current = selected_by_title.get(title_key)
        if current is None:
            selected_by_title[title_key] = item
            continue

        def candidate_score(candidate: ImportedItem) -> tuple[int, int, int]:
            candidate_check = checks[candidate.key()]
            candidate_title = (
                clean_amazon_title(candidate_check.get("title", ""))
                if candidate.source == "Amazon"
                else candidate_check.get("title", "")
            )
            similarity = title_similarity(candidate.title, candidate_title)
            return (
                int(bool(candidate_check.get("ok")) and bool(similarity["safe"])),
                int(bool(candidate_check.get("image"))),
                richness(candidate),
            )

        if candidate_score(item) > candidate_score(current):
            item.source_notes = list(dict.fromkeys(item.source_notes + current.source_notes))
            item.source_file = " | ".join(
                dict.fromkeys(
                    part.strip()
                    for part in f"{item.source_file}|{current.source_file}".split("|")
                )
            )
            selected_by_title[title_key] = item
        else:
            current.source_notes = list(dict.fromkeys(current.source_notes + item.source_notes))
            current.source_file = " | ".join(
                dict.fromkeys(
                    part.strip()
                    for part in f"{current.source_file}|{item.source_file}".split("|")
                )
            )
    merged = list(selected_by_title.values())

    by_link, by_title = existing_index(products)
    imported_products: list[dict[str, Any]] = []
    added = 0
    updated = 0
    active = 0
    manual = 0
    image_errors: list[dict[str, str]] = []

    for item in merged:
        check = checks[item.key()]
        checked_title = clean_amazon_title(check.get("title", "")) if item.source == "Amazon" else check.get("title", "")
        match = title_similarity(item.title, checked_title)
        image_relative = ""
        if check.get("ok") and match["safe"] and check.get("image"):
            filename = f"{slugify(item.title)}-{stable_product_id(item.product_id_prefix, item.title, item.link)[-8:]}.webp"
            image_relative = f"public/images/anuncios/{filename}"
            destination = root / image_relative
            if args.apply:
                try:
                    optimize_image_from_url(check["image"], destination)
                except Exception as error:  # noqa: BLE001
                    image_errors.append({"title": item.title, "error": str(error)})
                    image_relative = ""
        elif item.embedded_image_safe and item.embedded_image and match["safe"]:
            filename = f"{slugify(item.title)}-{stable_product_id('docx', item.title, item.link)[-8:]}.webp"
            image_relative = f"public/images/anuncios/{filename}"
            if args.apply:
                from io import BytesIO
                from PIL import Image

                destination = root / image_relative
                destination.parent.mkdir(parents=True, exist_ok=True)
                with Image.open(BytesIO(item.embedded_image)) as image:
                    if image.mode not in ("RGB", "RGBA"):
                        image = image.convert("RGB")
                    if image.mode == "RGBA":
                        background = Image.new("RGB", image.size, "white")
                        background.paste(image, mask=image.getchannel("A"))
                        image = background
                    image.thumbnail((1200, 1200), Image.Resampling.LANCZOS)
                    image.save(destination, format="WEBP", quality=82, method=6)

        existing = by_link.get(item.key()) or by_title.get(normalize_text(item.title))
        product = build_product(item, check, root, image_relative, existing)
        imported_products.append(product)
        if product["status"] == "ativo":
            active += 1
        else:
            manual += 1
        if existing:
            updated += 1
            products[products.index(existing)] = product
        else:
            added += 1
            products.append(product)

    store_created = ensure_toy_store(stores)
    rotation = (
        update_home_rotation(root, imported_products)
        if args.apply
        else {"prepared": len(imported_products), "active": active, "inactive": manual}
    )
    report = {
        "executedAt": today_iso(),
        "sourceFiles": extraction_counts,
        "rawExtracted": len(extracted),
        "uniqueAfterDeduplication": len(merged),
        "productsAdded": added,
        "productsUpdated": updated,
        "activeImported": active,
        "manualReviewImported": manual,
        "toyStoreCreated": store_created,
        "homeRotation": rotation,
        "imageErrors": image_errors,
        "manualReview": [
            {
                "id": product["id"],
                "title": product["name"],
                "link": product["affiliateLink"],
                "pendencias": product["pendencias"],
            }
            for product in imported_products
            if product["status"] != "ativo"
        ],
    }

    if args.apply:
        write_json(products_path, products)
        write_json(stores_path, stores)
        write_json(root / "dados" / "ultima-importacao-anuncios.json", report)
        write_json(root / "dados" / f"produtos-importados-{today_iso()}.json", imported_products)
        print(f"Aplicado: {added} adicionados, {updated} atualizados, {active} ativos, {manual} para revisão.")
    else:
        preview = root / "dados" / "preview-importacao-anuncios.json"
        write_json(preview, {"report": report, "products": imported_products})
        print(f"Prévia salva em {preview}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
