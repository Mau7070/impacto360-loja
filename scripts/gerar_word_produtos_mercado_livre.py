import argparse
import json
import re
import sys
from pathlib import Path
from urllib.parse import urlparse


DEFAULT_INPUT = Path("dados/mercado-livre-links.json")
DEFAULT_OUTPUT = Path("output")
USER_AGENT = "Impacto360Afiliado/1.0 (+https://impacto360afiliado.com.br)"

requests = None
Document = None
WD_BREAK = None
OxmlElement = None
qn = None
Inches = None


def carregar_dependencias():
    global requests, Document, WD_BREAK, OxmlElement, qn, Inches
    try:
        import requests as requests_mod
        from docx import Document as document_mod
        from docx.enum.text import WD_BREAK as wd_break_mod
        from docx.oxml import OxmlElement as oxml_element_mod
        from docx.oxml.ns import qn as qn_mod
        from docx.shared import Inches as inches_mod
    except ModuleNotFoundError as erro:
        pacote = erro.name or "dependencia"
        raise RuntimeError(
            f"Dependencia ausente: {pacote}. Instale com: "
            "pip install -r requirements-mercado-livre.txt"
        ) from erro

    requests = requests_mod
    Document = document_mod
    WD_BREAK = wd_break_mod
    OxmlElement = oxml_element_mod
    qn = qn_mod
    Inches = inches_mod


def normalizar_link(valor):
    return str(valor or "").strip()


def carregar_links(caminho):
    caminho = Path(caminho)
    if not caminho.exists():
        raise FileNotFoundError(f"Arquivo de links nao encontrado: {caminho}")

    if caminho.suffix.lower() == ".json":
        dados = json.loads(caminho.read_text(encoding="utf-8-sig"))
        links = []
        for item in dados:
            if isinstance(item, str):
                links.append(item)
            elif isinstance(item, dict):
                links.append(item.get("affiliateLink") or item.get("link") or item.get("url"))
        return deduplicar_links(links)

    texto = caminho.read_text(encoding="utf-8-sig")
    links = re.findall(r"https?://[^\s<>'\"]+", texto)
    return deduplicar_links(links)


def deduplicar_links(links):
    vistos = set()
    resultado = []
    for link in links:
        link = normalizar_link(link)
        if not link or link in vistos:
            continue
        if "meli.la" not in link and "mercadolivre.com" not in link:
            continue
        vistos.add(link)
        resultado.append(link)
    return resultado


def criar_sessao():
    sessao = requests.Session()
    sessao.headers.update({
        "User-Agent": USER_AGENT,
        "Accept": "application/json,text/html,*/*",
        "Accept-Language": "pt-BR,pt;q=0.9",
    })
    return sessao


def resolver_link_afiliado(sessao, url):
    resposta = sessao.get(url, allow_redirects=True, timeout=20)
    resposta.raise_for_status()
    return resposta.url


def extrair_item_id(url):
    padroes = [
        r"(MLB-?\d+)",
        r"/p/(MLB\d+)",
        r"itemId=(MLB\d+)",
    ]
    for padrao in padroes:
        achado = re.search(padrao, url, flags=re.IGNORECASE)
        if achado:
            return achado.group(1).replace("-", "").upper()
    raise ValueError("Nao encontrei o codigo MLB no link final.")


def buscar_item(sessao, item_id):
    resposta = sessao.get(f"https://api.mercadolibre.com/items/{item_id}", timeout=20)
    resposta.raise_for_status()
    return resposta.json()


def buscar_descricao(sessao, item_id):
    resposta = sessao.get(f"https://api.mercadolibre.com/items/{item_id}/description", timeout=20)
    if resposta.status_code == 200:
        return resposta.json().get("plain_text", "") or ""
    return ""


def resumir_descricao(texto, limite=650):
    texto = re.sub(r"\s+", " ", str(texto or "")).strip()
    if not texto:
        return "Informacao nao especificada pelo fornecedor."
    if len(texto) <= limite:
        return texto
    corte = texto[:limite].rsplit(" ", 1)[0]
    return f"{corte}..."


def moeda_brasil(valor):
    if valor is None:
        return "Consultar"
    try:
        return f"R$ {float(valor):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except (TypeError, ValueError):
        return str(valor)


def escolher_imagem(item):
    pictures = item.get("pictures") or []
    if pictures:
        return pictures[0].get("secure_url") or pictures[0].get("url")
    return item.get("secure_thumbnail") or item.get("thumbnail")


def extensao_imagem(url, content_type):
    parsed = urlparse(url)
    suffix = Path(parsed.path).suffix.lower()
    if suffix in {".jpg", ".jpeg", ".png", ".webp"}:
        return suffix
    if "png" in content_type:
        return ".png"
    if "webp" in content_type:
        return ".webp"
    return ".jpg"


def baixar_imagem(sessao, url, destino_base):
    if not url:
        return None
    resposta = sessao.get(url, timeout=25)
    resposta.raise_for_status()
    extensao = extensao_imagem(url, resposta.headers.get("Content-Type", ""))
    destino = destino_base.with_suffix(extensao)
    destino.write_bytes(resposta.content)
    return destino


def adicionar_hyperlink(paragrafo, texto, url):
    part = paragrafo.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    cor = OxmlElement("w:color")
    cor.set(qn("w:val"), "0563C1")
    rpr.append(cor)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)

    run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = texto
    run.append(text_node)
    hyperlink.append(run)
    paragrafo._p.append(hyperlink)


def produto_pendente(link_original, erro):
    return {
        "status": "pendente",
        "erro": str(erro),
        "titulo": "Produto Mercado Livre pendente",
        "preco": "Consultar",
        "descricao_resumida": "Informacao nao especificada pelo fornecedor.",
        "imagem_url": "",
        "imagem_local": "",
        "link_afiliado": link_original,
        "link_final": "",
        "item_id": "",
    }


def processar_link(sessao, link_original, indice, imagens_dir):
    try:
        link_final = resolver_link_afiliado(sessao, link_original)
        item_id = extrair_item_id(link_final)
        item = buscar_item(sessao, item_id)
        descricao = buscar_descricao(sessao, item_id)
        imagem_url = escolher_imagem(item)
        imagem_local = baixar_imagem(sessao, imagem_url, imagens_dir / f"produto-{indice:03d}")

        return {
            "status": "ativo",
            "erro": "",
            "titulo": item.get("title") or "Informacao nao especificada pelo fornecedor",
            "preco": moeda_brasil(item.get("price")),
            "descricao_resumida": resumir_descricao(descricao),
            "imagem_url": imagem_url or "",
            "imagem_local": str(imagem_local) if imagem_local else "",
            "link_afiliado": link_original,
            "link_final": link_final,
            "item_id": item_id,
            "condicao": item.get("condition") or "Informacao nao especificada pelo fornecedor",
            "permalink_publico": item.get("permalink") or "",
        }
    except Exception as erro:
        return produto_pendente(link_original, erro)


def gerar_word(produtos, output_dir):
    documento = Document()
    documento.add_heading("Produtos Mercado Livre - IMPACTO 360 AFILIADO", 0)
    documento.add_paragraph(
        "Relatorio gerado a partir dos links de afiliado originais. "
        "Os links clicaveis abaixo preservam o link de comissao enviado."
    )

    for indice, produto in enumerate(produtos, start=1):
        documento.add_heading(f"{indice}. {produto['titulo']}", level=1)

        imagem_local = produto.get("imagem_local")
        if imagem_local and Path(imagem_local).exists():
            try:
                documento.add_picture(imagem_local, width=Inches(3.0))
            except Exception:
                documento.add_paragraph("Imagem baixada, mas nao foi possivel inserir no Word.")
        else:
            documento.add_paragraph("Imagem: Informacao nao especificada pelo fornecedor.")

        documento.add_paragraph(f"Preco: {produto.get('preco') or 'Consultar'}")
        documento.add_paragraph(f"Status: {produto.get('status')}")
        if produto.get("item_id"):
            documento.add_paragraph(f"ID Mercado Livre: {produto.get('item_id')}")

        documento.add_paragraph("Descricao resumida:")
        documento.add_paragraph(produto.get("descricao_resumida") or "Informacao nao especificada pelo fornecedor.")

        paragrafo_link = documento.add_paragraph("Link de afiliado: ")
        adicionar_hyperlink(paragrafo_link, produto["link_afiliado"], produto["link_afiliado"])

        if produto.get("erro"):
            documento.add_paragraph(f"Erro/observacao: {produto['erro']}")

        if indice < len(produtos):
            documento.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    destino = output_dir / "produtos-mercado-livre.docx"
    documento.save(destino)
    return destino


def salvar_json(produtos, output_dir):
    destino = output_dir / "produtos-mercado-livre.json"
    destino.write_text(json.dumps(produtos, ensure_ascii=False, indent=2), encoding="utf-8")
    return destino


def main():
    parser = argparse.ArgumentParser(
        description="Gera Word com produtos Mercado Livre a partir de links de afiliado."
    )
    parser.add_argument("--links-file", default=str(DEFAULT_INPUT), help="Arquivo .json ou .txt com links.")
    parser.add_argument("--output-dir", default=str(DEFAULT_OUTPUT), help="Pasta de saida.")
    args = parser.parse_args()

    output_dir = Path(args.output_dir)
    imagens_dir = output_dir / "imagens"
    imagens_dir.mkdir(parents=True, exist_ok=True)

    links = carregar_links(args.links_file)
    if not links:
        print("Nenhum link Mercado Livre encontrado.", file=sys.stderr)
        return 1

    try:
        carregar_dependencias()
    except RuntimeError as erro:
        print(str(erro), file=sys.stderr)
        return 1

    sessao = criar_sessao()
    produtos = []
    for indice, link in enumerate(links, start=1):
        print(f"[{indice}/{len(links)}] Processando: {link}")
        produtos.append(processar_link(sessao, link, indice, imagens_dir))

    json_path = salvar_json(produtos, output_dir)
    word_path = gerar_word(produtos, output_dir)

    ativos = sum(1 for produto in produtos if produto["status"] == "ativo")
    pendentes = len(produtos) - ativos
    print(f"Concluido. Ativos: {ativos}. Pendentes: {pendentes}.")
    print(f"JSON: {json_path}")
    print(f"Word: {word_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
